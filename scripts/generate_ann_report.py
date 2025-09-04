import sys
from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import plotly.graph_objs as go
from io import BytesIO

# Local imports
BASE = Path(__file__).resolve().parents[1]
sys.path.append(str(BASE))
from src.app import get_ticker_data, prepare_df_numeric  # reuse data helpers
from src.features import add_basic_features
from src.labeling import create_reg_target, create_class_labels
from src.model import load_model, predict_with_model

MODELS_DIR = BASE / "models"
REPORTS_DIR = BASE / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

def find_latest_model(prefix: str, ticker: str, horizon: int | None = None):
    """Return latest matching model path for given prefix/ticker/horizon, else None."""
    patterns = []
    # examples: ann_reg_AAPL_h1.joblib or ann_class_AAPL_h1.joblib
    if horizon is not None:
        patterns.append(f"{prefix}{ticker}_h{horizon}.joblib")
        patterns.append(f"{prefix}{ticker}.joblib")
    else:
        patterns.append(f"{prefix}{ticker}.joblib")
    for name in patterns:
        p = MODELS_DIR / name
        if p.exists():
            return p
    # fallback: any file starting with prefix+ticker
    for p in sorted(MODELS_DIR.glob(f"{prefix}{ticker}*.joblib")):
        return p
    return None


def compute_reg_metrics(y_true: np.ndarray, y_pred: np.ndarray) -> dict:
    from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
    mae = float(mean_absolute_error(y_true, y_pred))
    rmse = float(np.sqrt(mean_squared_error(y_true, y_pred)))
    r2 = float(r2_score(y_true, y_pred))
    # MAPE with safe handling for zeros
    denom = np.where(y_true == 0, np.nan, np.abs(y_true))
    mape = float(np.nanmean(np.abs((y_true - y_pred) / denom)) * 100)
    # Directional accuracy
    dir_true = np.sign(np.diff(y_true))
    dir_pred = np.sign(np.diff(y_pred))
    dir_acc = float(np.mean(dir_true == dir_pred)) if len(dir_true) == len(dir_pred) and len(dir_true) > 0 else float('nan')
    return {"MAE": mae, "RMSE": rmse, "R2": r2, "MAPE%": mape, "DirAcc": dir_acc}


def line_chart_bytes(df: pd.DataFrame, title: str) -> bytes:
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=df['Date'], y=df['Actual'], name='Actual', mode='lines'))
    fig.add_trace(go.Scatter(x=df['Date'], y=df['Pred'], name='Pred', mode='lines'))
    fig.update_layout(title=title, xaxis_title='Date', yaxis_title='Price', template='plotly_white', legend=dict(orientation='h'))
    img_bytes = fig.to_image(format='png', scale=2)  # requires kaleido via plotly
    return img_bytes


def build_report(tickers: list[str], horizon: int = 1, rows: int = 600, out_name: str = None):
    doc = Document()
    doc.add_heading('StockMaster ANN Report', 0)
    doc.add_paragraph(f"Horizon: h{horizon} | Rows evaluated per asset: {rows}")
    doc.add_paragraph("Model: ANN (MLPRegressor and MLPClassifier pipelines as available)")
    doc.add_paragraph("")

    for tk in tickers:
        doc.add_heading(tk, level=1)
        # Load data from combined CSV
        df_raw = get_ticker_data(tk)
        if df_raw is None or df_raw.empty:
            doc.add_paragraph("No data found.")
            continue
        # Regression part
        reg_prefix = "ann_reg_"
        reg_model_path = find_latest_model(reg_prefix, tk, horizon=horizon)
        if reg_model_path and reg_model_path.exists():
            df_reg = create_reg_target(df_raw, horizon=horizon)
            df_reg = add_basic_features(df_reg)
            non_features = {"Date", "next_close"}
            features = [c for c in df_reg.columns if c not in non_features and np.issubdtype(df_reg[c].dtype, np.number)]
            # take last N rows
            df_eval = df_reg.tail(rows).copy()
            model = load_model(reg_model_path)
            preds = predict_with_model(model, df_eval, features, is_classifier=False)
            y_true = df_eval['next_close'].to_numpy()
            y_pred = np.asarray(preds)
            metrics = compute_reg_metrics(y_true, y_pred)
            # table
            doc.add_paragraph("Regression metrics:")
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            for i, k in enumerate(["MAE","RMSE","R2","MAPE%","DirAcc"]):
                hdr[i].text = k
            row = table.add_row().cells
            for i, k in enumerate(["MAE","RMSE","R2","MAPE%","DirAcc"]):
                row[i].text = f"{metrics[k]:.4f}" if isinstance(metrics[k], (int,float)) and not np.isnan(metrics[k]) else str(metrics[k])
            # chart
            plot_df = pd.DataFrame({"Date": df_eval['Date'], "Actual": y_true, "Pred": y_pred})
            try:
                img = line_chart_bytes(plot_df, f"{tk} ANN Regression h{horizon}")
                img_stream = BytesIO(img)
                doc.add_picture(img_stream, width=Inches(6.5))
            except Exception as e:
                doc.add_paragraph(f"Plot render failed: {e}")
        else:
            doc.add_paragraph("No ANN regressor model found.")
        # Classification part (optional)
        class_prefix = "ann_class_"
        class_model_path = find_latest_model(class_prefix, tk, horizon=horizon)
        if class_model_path and class_model_path.exists():
            df_cls = create_class_labels(df_raw, horizon=horizon)
            df_cls = add_basic_features(df_cls)
            non_features = {"Date", "next_close", "ret_next", "label"}
            features = [c for c in df_cls.columns if c not in non_features and np.issubdtype(df_cls[c].dtype, np.number)]
            df_eval = df_cls.tail(rows).copy()
            model = load_model(class_model_path)
            preds = predict_with_model(model, df_eval, features, is_classifier=True)
            y_true = df_eval['label'].to_numpy()
            y_pred = np.asarray(preds)
            from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score
            metrics = {
                "Accuracy": float(accuracy_score(y_true, y_pred)),
                "Precision": float(precision_score(y_true, y_pred, average='macro', zero_division=0)),
                "Recall": float(recall_score(y_true, y_pred, average='macro', zero_division=0)),
                "F1": float(f1_score(y_true, y_pred, average='macro', zero_division=0)),
            }
            doc.add_paragraph("Classification metrics:")
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            for i, k in enumerate(["Accuracy","Precision","Recall","F1"]):
                hdr[i].text = k
            row = table.add_row().cells
            for i, k in enumerate(["Accuracy","Precision","Recall","F1"]):
                row[i].text = f"{metrics[k]:.4f}"
        else:
            doc.add_paragraph("No ANN classifier model found.")

    out_name = out_name or f"ann_report_h{horizon}.docx"
    out_path = REPORTS_DIR / out_name
    doc.save(out_path)
    print(f"Report written: {out_path}")

if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument('--tickers', nargs='+', default=['AAPL','AMZN'])
    p.add_argument('--horizon', type=int, default=1)
    p.add_argument('--rows', type=int, default=600)
    p.add_argument('--out', type=str, default=None)
    args = p.parse_args()
    build_report(args.tickers, horizon=args.horizon, rows=args.rows, out_name=args.out)
